from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "literature_priority_notes.md"
DOCX_PATH = ROOT / "literature_priority_notes.docx"


def set_run_font(run, font_name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_style(paragraph, before=0, after=6, line_spacing=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_text_with_bold(paragraph, text, size=11):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("文献阅读整理")
    set_run_font(run, size=9, color=(100, 100, 100))


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_doc():
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, (46, 116, 181), 18, 10),
        ("Heading 2", 13, (46, 116, 181), 14, 7),
        ("Heading 3", 12, (31, 77, 120), 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    lines = text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_style(p, before=0, after=14, line_spacing=1.15)
            run = p.add_run(line[2:])
            set_run_font(run, size=20, bold=True, color=(11, 37, 69))
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue

        if line.startswith("- 论文"):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_style(p, before=6, after=4, line_spacing=1.25)
            run = p.add_run(line[2:])
            set_run_font(run, size=11, bold=True, color=(31, 77, 120))
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            label_match = re.match(r"- ([^：:]+[：:])\s*(.*)", stripped)
            p = doc.add_paragraph(style="List Bullet 2")
            set_paragraph_style(p, before=0, after=4, line_spacing=1.25)
            if label_match:
                label, rest = label_match.groups()
                r1 = p.add_run(label)
                set_run_font(r1, size=11, bold=True, color=(31, 77, 120))
                add_text_with_bold(p, rest, size=11)
            else:
                add_text_with_bold(p, stripped[2:], size=11)
            continue

        number_match = re.match(r"(\d+)\.\s+(.*)", line)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_style(p, before=0, after=4, line_spacing=1.25)
            add_text_with_bold(p, number_match.group(2), size=11)
            continue

        p = doc.add_paragraph()
        set_paragraph_style(p, before=0, after=6, line_spacing=1.25)
        add_text_with_bold(p, line, size=11)

    page_p = section.footer.paragraphs[0]
    page_p.add_run("  |  第 ")
    add_page_number(page_p)
    page_p.add_run(" 页")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
