from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "selected_literature_reference.docx"

CONTENT = r"""
第一优先级：核心必读文献

论文1：A new reinforcement learning-based variable speed limit control approach to improve traffic efficiency against freeway jam waves
出处：Transportation Research Part C: Emerging Technologies，2022。
主要内容：这篇论文研究高速公路拥堵波下的 RL-VSL 控制。它不是简单让智能体在仿真器里乱试，而是把控制过程分成“在线控制”和“离线训练”：在线阶段执行当前策略并收集交通状态、限速动作、控制效果；离线阶段再用这些数据更新 RL 策略。它还特别讨论了仿真模型和真实交通不一致的问题。
与我项目的关系：你的项目也是高速公路 VSL 控制，只是场景从“拥堵波”换成“事故瓶颈”。事故瓶颈会造成容量下降、速度突变和排队传播，所以这篇可以作为你 RL-VSL 主线文献。
可参考之处：这篇最值得你参考的是 RL-VSL 问题怎么搭起来。你可以看它怎样定义交通状态，比如上游路段的速度、密度、流量等；你的项目里可以把事故瓶颈上游分成几个检测区，把每个区的平均速度、车辆数、排队长度、速度波动作为状态。它的动作是 VSL 限速值；你的项目也可以把动作设成一组离散限速，例如 40、60、80、100 km/h，或者让安全主体和效率主体分别提出一个限速，再由规则或网络融合。它的奖励偏向交通效率，比如降低总耗费时间；你的项目不能只学这一点，而是可以把它作为效率主体的奖励设计参考，例如平均速度越高、通过事故瓶颈车辆越多、总延误越低，效率主体奖励越高。它讨论“仿真器训练有风险，因为仿真和真实不一致”；你的论文可以借这个点说明：事故瓶颈更复杂，单一 RL 策略容易学到偏仿真的策略，因此你提出双主体和自博弈式训练，希望让策略在安全与效率拉扯中更稳健。不能照搬的是：它主要目标是交通效率，不是安全-效率协商；你不能直接把它的奖励函数搬过来当总奖励，否则你的创新点会被削弱。

论文2：Optimal Coordination of Variable Speed Limits to Suppress Shock Waves
出处：IEEE Transactions on Intelligent Transportation Systems，2005。
主要内容：这篇论文用 MPC 控制 VSL 来抑制高速公路冲击波。它强调 VSL 不能随便跳变，因为限速突然下降会让驾驶员急刹，带来安全风险。所以它在优化里加入了安全约束，例如相邻限速变化不能太大，限速值也要符合现实标志牌的离散形式。
与我项目的关系：你的事故瓶颈场景会出现急减速、速度波动和冲击波传播，所以这篇是你写“为什么 VSL 要平滑、为什么要关注安全”的基础。
可参考之处：这篇最适合帮你设计安全主体。它说限速下降不能太突然，你可以把这个思想放进安全主体奖励里：如果当前限速和上一时刻限速差太大，就给惩罚；如果车辆出现急刹次数增加，也给惩罚。它把 VSL 设成离散限速，你的项目也可以采用离散动作空间，而不是连续输出任意速度。这样更符合实际，也更容易训练。它强调预测冲击波，你虽然不一定做 MPC，但可以借它的思想：状态不要只看事故点附近，也要看上游多个路段，因为事故影响会向上游传播。你还可以用它解释：如果只追求效率，智能体可能给出激进限速；如果只追求安全，可能过度降速。所以你的双主体框架就是把这两个倾向拆开，让它们“协商”。不能照搬的是 MPC 求解过程。你的项目是 RL，不需要复现它的优化器；你主要借它的安全约束、离散限速和冲击波控制思想。

论文3：Enhancing Transferability of Deep Reinforcement Learning-Based Variable Speed Limit Control Using Transfer Learning
出处：IEEE Transactions on Intelligent Transportation Systems，2021。
主要内容：这篇论文用 DDQN 做 VSL 控制，并研究一个场景中学到的策略能不能迁移到别的场景，例如超速、恶劣天气、通行能力下降。它发现相似场景之间迁移有效，但差异太大时可能效果不好。
与我项目的关系：你的事故瓶颈就是一种通行能力下降场景，而且事故严重程度可能不同，比如封一条车道、封两条车道、事故持续时间不同、交通需求不同。
可参考之处：这篇最适合你参考事故场景怎么设置多个难度版本。你可以像它一样设置不同目标场景：轻度事故、中度事故、重度事故；低流量、中流量、高流量；事故位置不同；事故持续时间不同。它用 DDQN 处理离散 VSL 动作，你如果目前算法还没完全定，可以把 DDQN 或 DQN 作为 baseline。你的双主体方法可以和“单智能体 DDQN-VSL”比较，证明双主体不是白加的。它关注 TTS、训练时间和不同场景效果，你可以参考它的实验写法：先在基础事故场景训练，再看方法在不同事故严重程度下表现是否稳定。它说场景相似性低时迁移可能失败，这对你有用：你可以在论文里说明事故瓶颈具有突发性和多样性，因此需要比普通单策略更稳健的训练机制。不能照搬的是迁移学习本身。如果你的项目不做迁移学习，不要硬说自己用了迁移；可以说它启发你设计多事故场景测试和泛化实验。

论文4：Integrated Feedback Ramp Metering and Mainstream Traffic Flow Control on Motorways Using Variable Speed Limits
出处：Transportation Research Part C: Emerging Technologies，2014。
主要内容：这篇论文把匝道流量控制和 VSL 主线控制结合起来。它认为匝道控制会受匝道排队空间限制，当匝道排队太长时，VSL 可以继续在主线上游调节交通流。
与我项目的关系：你的方法不一定控制匝道，但同样是用 VSL 在瓶颈上游干预主线交通。
可参考之处：这篇可以帮你解释 VSL 为什么适合做事故瓶颈上游控制。它把 VSL 当作主线交通流控制手段，而不只是安全提示牌。你的项目可以沿用这个思想：事故发生后，不能直接改变事故点通行能力，但可以通过上游限速调节车辆到达事故点的速度和密度。它的反馈控制思想也可以变成你的 baseline：比如设置一个简单规则，当事故上游密度过高时降低限速，当速度恢复时提高限速。这样你可以拿自己的 RL 方法和传统反馈规则比较。它还说明控制周期问题，RM 和 VSL 可能不同步。你的项目也要注意：限速不能每秒乱变，最好设置固定控制间隔，例如每 30 秒或 60 秒更新一次 VSL。不能照搬的是它的 RM-VSL 集成框架，因为你的研究重点不是匝道排队，而是事故瓶颈安全-效率交互。

第二优先级：重要参考文献

论文5：A physics-informed reinforcement learning-based strategy for local and coordinated ramp metering
出处：Transportation Research Part C: Emerging Technologies，2022。
主要内容：这篇论文用物理信息 RL 做匝道控制。它把历史数据和交通流模型生成的数据结合起来训练 RL，避免智能体完全依赖随机探索。
与我项目的关系：它不是 VSL，但和你一样面对 RL 交通控制中的一个现实问题：真实道路不能让智能体随便试错。
可参考之处：这篇可以帮你改进训练方式和论文论证。如果你只在 SUMO 中训练，可以参考它的说法：交通控制中的 RL 随机探索在现实中有风险，因此需要先在仿真中训练，再通过更稳健的机制减少不合理动作。它把交通流模型生成的数据叫作合成数据，你也可以在实验里设计不同事故场景生成训练数据，让智能体见过更多事故扰动。它把协调控制问题拆成子系统，你的项目也可以借这个思想：如果你有多个 VSL 控制点，可以不要让一个巨大智能体一次控制所有限速，而是让不同主体/不同区域分别提出控制意见。它比较了表格型 RL 和深度 RL，并发现数据有限时深度 RL 不一定更好。这个提醒很重要：你的实验不要只追求模型复杂，最好设置简单 baseline，比如固定限速、反馈限速、单智能体 DQN/PPO。不能照搬的是它的匝道控制动作，因为你的动作是限速，不是匝道放行率。

论文6：Expert Level Control of Ramp Metering Based on Multi-Task Deep Reinforcement Learning
出处：IEEE Transactions on Intelligent Transportation Systems，2018。
主要内容：这篇论文用深度 RL 做匝道控制，并提出 MWR，让多个智能体之间共享经验，同时允许每个智能体适应自己所在位置的局部交通状态。
与我项目的关系：你的安全主体和效率主体不是控制不同匝道，但它们也是两个有不同偏好的学习主体。
可参考之处：这篇可以帮你理解两个主体不是简单复制，而是要有差异化职责。它的方法思想是：多个智能体可以共享一部分知识，但不能完全一样，因为每个位置的交通状态不同。放到你的项目里，可以理解成：安全主体和效率主体可以共享同一个交通状态输入，但奖励函数和输出倾向不同。安全主体可以更关注急刹、速度波动、限速跳变；效率主体可以更关注平均速度、通过量、延误。如果你后面写方法，可以说你的双主体不是为了增加复杂度，而是为了让不同目标形成不同策略偏好，再通过融合机制得到最终限速。它也适合启发你做消融实验：去掉安全主体，只保留效率主体；去掉效率主体，只保留安全主体；两个主体都保留，看是否安全和效率更平衡。不能照搬的是 MWR 公式本身，除非你真的实现多任务权重正则化。

论文7：Grandmaster level in StarCraft II using multi-agent reinforcement learning
出处：Nature，2019。
主要内容：这篇论文用 league training 训练 AlphaStar。它发现简单 self-play 可能只适应当前对手，容易策略循环，所以加入历史策略和针对性反策略，让训练对象更多样。
与我项目的关系：你的“自博弈启发”如果只写 AlphaGo Zero 会有点单薄，AlphaStar 可以帮你解释复杂多主体环境中为什么需要更稳健的交互训练。
可参考之处：这篇可以启发你的自博弈训练流程。你可以不把安全主体和效率主体看成真正敌人，而是看成两个偏好不同的“提案者”。训练时可以保存历史版本的安全主体和效率主体，让当前主体不只和最新版对方交互，也和历史版本交互。这样可以减少策略只适应当前对方的情况。举个简单版本：每隔若干轮保存一次安全主体和效率主体参数；训练当前效率主体时，随机选择一个历史安全主体作为交互对象；训练当前安全主体时，也随机选择历史效率主体。这样写起来就比“我用了自博弈”更具体，也更像一个真的方法。不能照搬的是游戏里的胜负机制。你的安全和效率不是零和关系，不能写成一个赢一个输；应该写成“差异化目标之间的策略交互”。

论文8：Mastering the Game of Go without Human Knowledge
出处：Nature，2017。
主要内容：这篇论文提出 AlphaGo Zero，通过自我对弈生成训练数据，在没有人类示范的情况下不断提升策略。
与我项目的关系：它是 self-play 的重要思想来源，但你的项目不是棋类游戏，所以只能作为方法启发。
可参考之处：这篇可以帮你解释为什么要让策略通过反复交互改进。放到你的项目里，可以写成：事故瓶颈下安全和效率的折中很难人工预设固定权重，因此让两个主体在反复交互中学习不同状态下的让步关系。例如：事故刚发生、上游速度波动大时，安全主体提案可能更低；当排队过长、平均速度过低时，效率主体提案可能更高。通过反复训练，融合机制学习什么时候偏安全、什么时候偏效率。如果你要实现一个简单版本，可以把每轮训练后的策略保存下来，让新策略与旧策略生成的数据共同训练，形成“自我迭代”的味道。不能照搬的是 AlphaGo Zero 的棋盘搜索和胜负奖励。你的交通任务没有天然胜负，必须用交通指标构造奖励。

第三优先级：背景补充文献

论文9：Traffic-Responsive Linked Ramp-Metering Control
出处：IEEE Transactions on Intelligent Transportation Systems，2008。
主要内容：这篇论文提出 HERO 联动匝道控制，核心是多个连续匝道不能只靠各自局部控制，而要根据上下游状态联动。
与我项目的关系：你的事故瓶颈上游也不是一个点的问题，排队和速度扰动会传播，所以多个路段状态都要看。
可参考之处：这篇可以启发你的状态空间设计。不要只把事故点附近速度作为状态，可以加入事故点上游多个检测段的速度、密度、车辆数、排队长度。它说明局部控制可能在复杂场景中误用，你可以据此说明：如果只根据一个检测点限速，可能导致上游新拥堵或限速过度；所以你的方法需要考虑多个路段的整体状态。它还可以作为传统协调控制 baseline 的背景，但不建议花太多篇幅。不能照搬的是 HERO 匝道控制规则，因为它控制的是入口匝道流量，不是主线 VSL。

论文10：Operational Characteristics of Mixed-Autonomy Traffic Flow on the Freeway With On- and Off-Ramps and Weaving Sections: An RL-Based Approach
出处：IEEE Transactions on Intelligent Transportation Systems，2022。
主要内容：这篇论文用 SUMO 研究 RL 自动驾驶策略对高速公路混合交通流的影响，评价拥堵、安全、效率和舒适性。
与我项目的关系：它不是 VSL，但它的实验评价维度与你很像。
可参考之处：这篇最适合你参考评价指标怎么写得丰富。它不只看平均速度，还看拥堵模式、安全性和舒适性。你的项目也可以避免只看一个平均速度指标。安全方面可以看急刹次数、速度标准差、最大减速度、冲突风险；效率方面可以看平均速度、通过车辆数、总延误、排队长度；平滑性方面可以看限速变化次数和限速变化幅度。它还提醒你：RL 策略可能缓解拥堵，但也可能牺牲舒适性。对应到你的项目，就是效率主体可能提高平均速度，但导致急刹或速度波动；安全主体可能降低急刹，但造成低速和排队。所以这篇能支撑你做双目标分析。不能照搬的是 AV 驾驶策略，因为你的控制对象不是车辆加速度或换道，而是 VSL 标志限速。

论文11：CoTV: Cooperative Control for Traffic Light Signals and Connected Autonomous Vehicles Using Deep Reinforcement Learning
出处：IEEE Transactions on Intelligent Transportation Systems，2023。
主要内容：这篇论文用 MARL 协同控制交通信号和 CAV，强调交通控制不能只看出行时间，还要考虑能耗、排放、安全和部署成本。
与我项目的关系：它的场景是城市交叉口，不是高速事故瓶颈，但“异质智能体协作”和“多目标交通控制”对你有启发。
可参考之处：这篇可以启发你的双主体信息交互方式。CoTV 不是让所有智能体互相传一大堆信息，而是只交换必要状态。你的项目也可以这样设计：安全主体和效率主体共享同一组交通状态，但输出不同限速提案；融合模块只接收两个提案和关键状态，不需要复杂通信。它还可以帮你写多目标动机：未来交通控制不应该只优化通行时间，你的事故瓶颈控制也不应只优化平均速度，而要兼顾急刹、速度波动和平滑限速。如果你使用 PPO 或 MARL 框架，它也能作为多智能体交通控制的参考文献。不能照搬的是交通信号和 CAV 协同结构，因为你的两个主体是目标偏好不同，不是物理设备不同。

第四优先级：低相关或可选文献

论文12：Deep Reinforcement Learning for Autonomous Driving: A Survey
出处：IEEE Transactions on Intelligent Transportation Systems，2022。
主要内容：这篇综述介绍 DRL 在自动驾驶中的应用，包括驾驶策略、规划、控制、仿真训练和真实部署挑战。
与我项目的关系：它不是 VSL 论文，也不是事故瓶颈论文，只能作为大背景。
可参考之处：这篇主要用来写泛泛的 DRL 背景。如果你需要解释什么是 DRL、为什么交通任务适合建模为序贯决策问题，可以引用它。如果你想说明仿真训练和真实部署之间存在差距，也可以引用它。但它不能支撑你的核心创新，因为它没有研究 VSL、事故瓶颈、安全效率双主体。
不能重点写：不要在文献综述里花太多篇幅介绍自动驾驶 DRL，否则会偏离你的项目主题。

你真正可以照着用的项目设计清单

状态空间可以参考 RL-VSL 和协调控制文献：事故点上游多个路段的平均速度、密度、流量、车辆数、排队长度、速度标准差、上一时刻限速。
动作空间可以参考 VSL 文献：离散限速值，例如 40/60/80/100 km/h，或者每个主体各自提出一个离散限速。
安全主体奖励可以参考 Hegyi 2005 和混合交通流论文：惩罚急刹次数、速度波动、限速骤降、过大减速度。
效率主体奖励可以参考 Han 2022 和 Ke 2021：奖励平均速度、通过车辆数，惩罚总延误、TTS、排队长度。
融合机制可以写成：最终限速由安全提案和效率提案共同决定，例如加权融合、取较保守值、或由一个协调模块根据交通状态动态选择权重。
自博弈部分可以参考 AlphaGo Zero 和 AlphaStar：保存历史策略，让当前安全主体/效率主体与历史版本交互，避免只适应当前对方。
Baseline 可以设置：无控制、固定限速、反馈式 VSL、单智能体 DQN/PPO-VSL、只有安全主体、只有效率主体、双主体完整方法。
评价指标最好分三组：效率指标、安全指标、控制平滑指标。这样你的“安全-效率交互”才看得出来。
"""


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para(p, before=0, after=6, line=1.25):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_labeled_paragraph(doc, label, body):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=5)
    r = p.add_run(label + "：")
    set_font(r, bold=True, color=(31, 77, 120))
    r = p.add_run(body)
    set_font(r)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, (46, 116, 181)),
        ("Heading 2", 13, (46, 116, 181)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(title, after=14, line=1.15)
    r = title.add_run("文献优先级与项目参考说明")
    set_font(r, size=20, bold=True, color=(11, 37, 69))

    current_paper = None
    for raw in CONTENT.strip().splitlines():
        line = raw.strip()
        if not line:
            continue
        if re.match(r"^(第一|第二|第三|第四)优先级", line):
            doc.add_heading(line, level=1)
        elif line.startswith("论文"):
            current_paper = line
            p = doc.add_paragraph()
            set_para(p, before=6, after=4)
            r = p.add_run(line)
            set_font(r, bold=True, color=(31, 77, 120))
        elif line.startswith("你真正可以照着用"):
            doc.add_heading(line, level=1)
        elif "：" in line and line.split("：", 1)[0] in {"出处", "主要内容", "与我项目的关系", "可参考之处", "不能重点写"}:
            label, body = line.split("：", 1)
            add_labeled_paragraph(doc, label, body)
        elif current_paper is None or line.startswith(("状态空间", "动作空间", "安全主体", "效率主体", "融合机制", "自博弈", "Baseline", "评价指标")):
            p = doc.add_paragraph(style="List Bullet")
            set_para(p, after=4)
            r = p.add_run(line)
            set_font(r)
        else:
            p = doc.add_paragraph()
            set_para(p)
            r = p.add_run(line)
            set_font(r)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
